#!/usr/bin/env python3
"""
Generate Technical Description (Part B) for Sustainability Signals.
EU proposal formatting: Times New Roman 11pt, A4, 15mm margins, single spacing.
Maximum 12 pages including cover page.
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "SustainabilitySignals_TechnicalDescription.docx")

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_paragraph(doc, text, style=None, bold=False, font_size=11, alignment=None, space_after=Pt(4), space_before=Pt(0), font_name="Times New Roman", color=None, italic=False):
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if alignment:
        p.alignment = alignment
    # Set line spacing to single
    p.paragraph_format.line_spacing = Pt(13)
    return p

def add_mixed_paragraph(doc, runs_data, space_after=Pt(4), space_before=Pt(0), alignment=None):
    """Add paragraph with mixed formatting. runs_data is list of (text, bold, italic, font_size, color)."""
    p = doc.add_paragraph()
    for text, bold, italic, font_size, color in runs_data:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(13)
    if alignment:
        p.alignment = alignment
    return p

def add_heading_styled(doc, text, level=1, font_size=14, color=None, space_before=Pt(12), space_after=Pt(6)):
    """Add heading with proper formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
    h.paragraph_format.space_before = space_before
    h.paragraph_format.space_after = space_after
    h.paragraph_format.line_spacing = Pt(14)
    return h

def add_bullet(doc, text, bold_prefix="", font_size=11, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.style = "List Bullet"
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = "Times New Roman"
        run_b.font.size = Pt(font_size)
        run_b.font.bold = True
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = Pt(12)
        set_cell_shading(cell, "1B5E20")
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = Pt(12)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8F5E9")
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    
    return table


def build_document():
    doc = Document()
    
    # --- Page Setup: A4, 15mm margins ---
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.line_spacing = Pt(13)
    
    # =========================================================================
    # COVER PAGE (Page 1)
    # =========================================================================
    add_paragraph(doc, "", font_size=11, space_after=Pt(40))
    add_paragraph(doc, "Technical Description (Part B)", bold=True, font_size=20,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20),
                  color=RGBColor(0x1B, 0x5E, 0x20))
    add_paragraph(doc, "", font_size=11, space_after=Pt(10))
    add_paragraph(doc, "SUSTAINABILITY SIGNALS", bold=True, font_size=18,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8),
                  color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph(doc, "An Open, AI-Powered Platform for Transparent ESG Disclosure Quality Assessment", 
                  font_size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30),
                  italic=True, color=RGBColor(0x42, 0x42, 0x42))
    
    add_paragraph(doc, "", font_size=11, space_after=Pt(20))
    
    # Cover metadata
    cover_items = [
        ("Call Identifier:", "HORIZON-CL5-2026-D4-01"),
        ("Topic:", "Sustainable Finance — Transparent ESG Assessment Tools"),
        ("Type of Action:", "Innovation Action (IA)"),
        ("Proposal Full Title:", "Sustainability Signals: Open AI-Driven Disclosure Quality and ESG Entity Extraction Platform"),
        ("Acronym:", "SUSTAINABILITY-SIGNALS"),
        ("Duration:", "36 months"),
        ("Version:", "1.0"),
        ("Date:", "February 2026"),
    ]
    for label, value in cover_items:
        add_mixed_paragraph(doc, [
            (label + " ", True, False, 11, RGBColor(0x33, 0x33, 0x33)),
            (value, False, False, 11, None),
        ], space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_paragraph(doc, "", font_size=11, space_after=Pt(40))
    add_paragraph(doc, "This proposal is a self-contained document.", font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
                  color=RGBColor(0x75, 0x75, 0x75))
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1: EXCELLENCE
    # =========================================================================
    add_heading_styled(doc, "1. Excellence", level=1, font_size=15, 
                       color=RGBColor(0x1B, 0x5E, 0x20), space_before=Pt(6))
    
    # --- 1.1 Objectives ---
    add_heading_styled(doc, "1.1 Objectives and Ambition", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc, 
        "Sustainability Signals is an open-source, AI-powered platform that brings transparency and rigour "
        "to Environmental, Social, and Governance (ESG) disclosure quality assessment. The platform addresses "
        "a critical gap in the sustainable finance ecosystem: the lack of standardised, reproducible, and "
        "evidence-based tools for evaluating the quality of corporate sustainability disclosures.")
    
    add_paragraph(doc,
        "The principal objectives are:", bold=True, space_after=Pt(2))
    
    objectives = [
        ("O1 — Automated Disclosure Quality Scoring: ", 
         "Develop and deploy a deterministic, regex-based scoring engine (method: regex-v4.1) that evaluates "
         "sustainability reports across four auditable subscores — Completeness (35%), Consistency (25%), "
         "Assurance (20%), and Transparency (20%) — with full evidence traceability."),
        ("O2 — ESG Entity Extraction: ",
         "Implement a hybrid NLP pipeline combining FinBERT-ESG-9 classification with LLM-based structured "
         "entity extraction (Llama 3.3 70B) to identify and categorise ESG-relevant entities across 10 classes: "
         "GHG emissions, climate targets, energy, water, waste, biodiversity, social metrics, governance policies, "
         "financial ESG, and regulatory indicators."),
        ("O3 — Grounded AI-Assisted Analysis: ",
         "Deploy a retrieval-augmented generation (RAG) system for report-grounded question-answering, enabling "
         "analysts to interrogate sustainability reports with citations anchored to source text."),
        ("O4 — Open & Reproducible Methodology: ",
         "Publish all scoring algorithms, feature detection rules, and subscore formulas as open-source code, "
         "enabling independent verification and academic scrutiny of ESG assessments."),
        ("O5 — Scalable Coverage Universe: ",
         "Build a comprehensive, searchable database of sustainability reports with in-browser PDF viewing, "
         "covering multiple GICS sectors, geographies, and reporting years."),
    ]
    for prefix, text in objectives:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)
    
    add_paragraph(doc, "", space_after=Pt(2))
    add_paragraph(doc,
        "These objectives directly respond to the European Commission's Sustainable Finance Strategy and the "
        "Corporate Sustainability Reporting Directive (CSRD), which mandates enhanced disclosure standards but "
        "provides limited tooling for systematic disclosure quality assessment.", space_after=Pt(6))

    # --- 1.2 Relation to the Work Programme ---
    add_heading_styled(doc, "1.2 Relation to the Work Programme", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "This proposal aligns with Horizon Europe Cluster 5 (Climate, Energy and Mobility) and Cluster 6 "
        "(Food, Bioeconomy, Natural Resources, Agriculture, and Environment), specifically addressing the need "
        "for digital tools that enhance the quality and comparability of sustainability information. "
        "The platform's alignment with ESRS, GRI, SASB, TCFD, ISSB (IFRS S1/S2), EU Taxonomy, TNFD, "
        "and CDP frameworks ensures comprehensive coverage of the evolving European regulatory landscape.")
    
    add_paragraph(doc,
        "The project also contributes to the European Green Deal's data infrastructure objectives by creating "
        "publicly accessible tools for assessing corporate environmental and social commitments, directly "
        "supporting the EU Taxonomy Regulation's transparency requirements.")

    # --- 1.3 Concept and Methodology ---
    add_heading_styled(doc, "1.3 Concept and Methodology", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The Sustainability Signals platform implements a multi-layered analytical pipeline that transforms "
        "raw sustainability report PDFs into structured, scored, and searchable intelligence. The architecture "
        "is designed around three core principles: transparency, reproducibility, and scalability.",
        space_after=Pt(4))
    
    add_paragraph(doc, "1.3.1 Technical Architecture", bold=True, font_size=11, space_after=Pt(2))
    
    # Architecture table
    create_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 19, TypeScript, Vite 7, Tailwind CSS 4", "Interactive UI with in-browser PDF viewing"],
            ["Edge API", "Cloudflare Pages Functions (serverless)", "Stateless API endpoints at edge locations"],
            ["Storage", "Cloudflare R2 (S3-compatible)", "PDFs, cached markdown, DQ scores, entity results"],
            ["AI Runtime", "Cloudflare Workers AI", "LLM inference (chat, evidence refinement, extraction)"],
            ["Search", "Cloudflare Vectorize", "Semantic retrieval for RAG-based Q&A"],
            ["ML Model", "FinBERT-ESG-9 (PyTorch)", "9-category ESG text classification"],
            ["Container", "Docker / Azure Container Apps", "FinBERT inference server (Flask)"],
        ],
        col_widths=[Cm(2.5), Cm(6.5), Cm(8.5)]
    )
    
    add_paragraph(doc, "", space_after=Pt(4))
    add_paragraph(doc, "1.3.2 Disclosure Quality (DQ) Scoring Engine", bold=True, font_size=11, space_after=Pt(2))
    
    add_paragraph(doc,
        "The DQ engine is the core innovation of the platform. It implements a deterministic, regex-based "
        "feature detection system (method kind: regex-v4.1) operating over normalised report text. "
        "The pipeline proceeds as follows:")
    
    dq_steps = [
        ("PDF Ingestion: ", "Reports are converted from PDF to markdown using AI-assisted extraction "
         "(Workers AI toMarkdown API). The resulting text undergoes normalisation: removal of conversion noise, "
         "markdown artefact cleanup, zero-width character removal, and unicode normalisation."),
        ("Text Preparation: ", "Large documents (>800,000 characters) are sampled using a head-tail strategy "
         "(420K head + remaining tail). Page segmentation identifies page boundaries. Boilerplate lines "
         "(repeated headers/footers appearing in ≥25% of pages) are detected and excluded."),
        ("Feature Detection: ", "Over 80 disclosure features are detected across 12 feature families using "
         "calibrated regular expressions. Each feature returns: boolean presence, evidence quotes (top 3 ranked), "
         "occurrence depth, and distinct page coverage."),
        ("Evidence Ranking: ", "Quotes are ranked by a quality function favouring: percentage presence, "
         "large numbers, ESG-specific units (tCO₂e, MWh, kWh, GJ, tonnes), year mentions, target language, "
         "scope mentions, table-like rows, and page references."),
        ("Subscore Computation: ", "Four subscores are computed from weighted feature combinations: "
         "Completeness (breadth of topic coverage), Consistency (methodological rigour and data quality), "
         "Assurance (external verification level), and Transparency (openness and forward-looking clarity)."),
        ("Overall Score: ", "The weighted composite score = 0.35×Completeness + 0.25×Consistency + "
         "0.20×Assurance + 0.20×Transparency, banded as High (≥75), Medium (50–74), or Low (<50)."),
        ("Evidence Refinement: ", "Optionally, extracted evidence quotes are cleaned using Llama 3.1 8B "
         "(temperature=0.1) to remove PDF artefacts while preserving all numerical data and facts."),
    ]
    for prefix, text in dq_steps:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)
    
    add_paragraph(doc, "", space_after=Pt(3))
    add_paragraph(doc, "1.3.3 Feature Families", bold=True, font_size=11, space_after=Pt(2))
    
    add_paragraph(doc,
        "The engine detects features across the following families, aligned with major ESG reporting standards:")
    
    create_table(doc,
        ["Family", "Key Features", "Standards Alignment"],
        [
            ["Frameworks & Standards", "ESRS, CSRD, GRI, SASB, TCFD, ISSB, EU Taxonomy, TNFD, CDP, SDGs, Paris", "Direct framework detection"],
            ["Materiality", "Double materiality, assessment, matrix, IRO, value chain", "CSRD Article 19a, ESRS 1"],
            ["Governance", "Board oversight, audit committee, ESG remuneration, whistleblower, data privacy", "GRI 2-9 to 2-26"],
            ["Climate & Emissions", "Scope 1/2/3, GHG Protocol, base year, emissions intensity, Scope 3 categories", "GHG Protocol, IFRS S2"],
            ["Targets", "Net zero, SBTi, transition plan, quantitative/interim targets, progress tracking", "Paris Agreement, TCFD"],
            ["Environment", "Energy, water, waste, renewables, biodiversity, circular economy, pollution", "GRI 300 series, TNFD"],
            ["Social", "Workforce, safety, diversity, human rights, living wage, just transition, turnover", "GRI 400 series, ESRS S1-S4"],
            ["Assurance", "Limited/reasonable assurance, ISAE 3000, AA1000, named provider, scope", "ISAE 3000, CSRD Art. 34"],
        ],
        col_widths=[Cm(3.2), Cm(8.0), Cm(5.5)]
    )
    
    add_paragraph(doc, "", space_after=Pt(3))
    add_paragraph(doc, "1.3.4 Subscore Formulas", bold=True, font_size=11, space_after=Pt(2))
    
    create_table(doc,
        ["Subscore", "Weight", "Key Components (Maximum Points)", "Focus"],
        [
            ["Completeness", "35%", "Frameworks (18), Climate/Emissions (22), Social (12), Materiality (10), Governance (8), Targets (10), Environment (8), EU Taxonomy (4), Sector-specific (2)", "Breadth and depth of reporting"],
            ["Consistency", "25%", "Methodology/boundary/data quality (20), Comparative years (18), Quant. density (15), GHG Protocol (10), Limitations (10), Control environment (8), Reporting period (5), Targets (5)", "Methodological rigour"],
            ["Assurance", "20%", "Base level: None(0) → Limited(60) → Reasonable(82) → Both(90). Bonuses: Standard(+6), Named provider(+5), Scope(+4), Breadth(+3)", "External verification"],
            ["Transparency", "20%", "Forward-looking (15), Limitations (15), Methodology (10), Stakeholder engagement (8), Scope 3 detail (7), Targets+progress (8), ESRS datapoints (5), Financial connectivity (5)", "Openness and clarity"],
        ],
        col_widths=[Cm(2.2), Cm(1.3), Cm(10.0), Cm(3.0)]
    )
    
    add_paragraph(doc, "", space_after=Pt(3))
    add_paragraph(doc, "1.3.5 Entity Extraction Pipeline", bold=True, font_size=11, space_after=Pt(2))
    
    add_paragraph(doc,
        "The entity extraction system (method: hybrid-finbert9-langextract-v2) implements a cost-efficient "
        "two-stage pipeline that combines specialised classification with large language model extraction:")
    
    ee_steps = [
        ("Stage 1 — FinBERT-ESG-9 Routing: ",
         "Report text is chunked (12K chars, 600 overlap) and classified by a fine-tuned FinBERT model into "
         "9 ESG categories (Climate Change, Natural Capital, Pollution & Waste, Human Capital, Product Liability, "
         "Community Relations, Corporate Governance, Business Ethics & Values, Non-ESG). Non-ESG chunks are "
         "filtered out, yielding 40–60% reduction in downstream LLM inference cost."),
        ("Stage 2 — LLM Entity Extraction: ",
         "ESG-relevant chunks are processed by Llama 3.3 70B with few-shot prompting to extract structured "
         "entities across 10 classes: ghg_emissions, climate_target, energy, water, waste, biodiversity, "
         "social_metric, governance_policy, financial_esg, and regulatory."),
        ("Pillar Mapping: ",
         "Extracted entities inherit E/S/G pillar classification from the FinBERT routing stage, enabling "
         "pillar-level aggregation and visualisation."),
    ]
    for prefix, text in ee_steps:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)
    
    add_paragraph(doc, "", space_after=Pt(3))
    add_paragraph(doc, "1.3.6 Retrieval-Augmented Generation (RAG) Chat", bold=True, font_size=11, space_after=Pt(2))
    
    add_paragraph(doc,
        "The platform provides report-grounded AI question-answering via a RAG architecture. The client "
        "extracts page text using PDF.js and sends it as context (capped at 120,000 characters). When Vectorize "
        "is enabled, the system performs semantic retrieval over pre-indexed report chunks. The LLM (configurable; "
        "default: Mistral Small 3.1 24B) is instructed to answer only from provided report text, never hallucinate "
        "facts, and cite specific page numbers. Background ingestion converts new PDFs to markdown, chunks them, "
        "and upserts embeddings to Vectorize using waitUntil() for non-blocking operation.")

    # --- 1.4 Novelty ---
    add_heading_styled(doc, "1.4 Novelty and Originality", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "Sustainability Signals introduces several novel contributions beyond the current state-of-the-art:")
    
    novelty_items = [
        ("Deterministic + Explainable DQ Scoring: ",
         "Unlike black-box ESG ratings from commercial providers, the regex-v4.1 engine produces fully auditable "
         "scores with traceable evidence quotes, page references, and feature-level breakdowns. Every point in "
         "the scoring formula can be independently verified."),
        ("Hybrid FinBERT + LLM Entity Extraction: ",
         "The two-stage pipeline is novel in combining a lightweight, specialised ESG classifier for content "
         "routing with a powerful generalist LLM for structured extraction, achieving cost efficiency without "
         "sacrificing extraction quality."),
        ("Edge-Native Architecture: ",
         "The entire platform runs on Cloudflare's edge network, with serverless API endpoints, R2 object "
         "storage, Workers AI inference, and Vectorize search — all at edge latency. This eliminates the need "
         "for traditional server infrastructure."),
        ("Full Open-Source Transparency: ",
         "All scoring algorithms, feature detection patterns, subscore formulas, and evidence ranking logic are "
         "published as open-source code, enabling academic scrutiny, regulatory auditing, and community-driven "
         "improvement — a first for ESG disclosure quality assessment tools."),
    ]
    for prefix, text in novelty_items:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)

    # =========================================================================
    # SECTION 2: IMPACT
    # =========================================================================
    add_heading_styled(doc, "2. Impact", level=1, font_size=15,
                       color=RGBColor(0x1B, 0x5E, 0x20), space_before=Pt(14))
    
    # --- 2.1 Expected Impact ---
    add_heading_styled(doc, "2.1 Expected Impacts", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The platform is expected to generate significant impact across multiple dimensions of the sustainable "
        "finance ecosystem:")
    
    impacts = [
        ("Regulatory Alignment: ",
         "By automating CSRD-aligned disclosure quality assessment and supporting ESRS, GRI, TCFD, and ISSB "
         "framework detection, the platform reduces compliance costs for reporting entities and provides regulators "
         "with scalable oversight tools."),
        ("Market Transparency: ",
         "The open, reproducible scoring methodology addresses greenwashing concerns by enabling investors, "
         "analysts, and civil society to independently verify ESG disclosure quality claims — a capability "
         "currently unavailable from commercial ESG rating providers."),
        ("Academic Research: ",
         "The fully documented scoring methodology (DQ_MECHANISM.md) and open-source codebase create a "
         "reproducible benchmark for sustainability disclosure research, enabling comparative studies across "
         "sectors, geographies, and reporting years."),
        ("SME Accessibility: ",
         "The edge-native, serverless architecture delivers enterprise-grade ESG analysis at near-zero marginal "
         "cost, making professional-quality disclosure assessment accessible to small and medium enterprises "
         "preparing for CSRD compliance."),
        ("Environmental Benefit: ",
         "The edge-computing architecture minimises energy consumption compared to traditional cloud deployments, "
         "while the FinBERT routing stage reduces LLM inference costs by 40–60%, directly lowering the carbon "
         "footprint of AI-powered ESG analysis."),
    ]
    for prefix, text in impacts:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)

    # --- 2.2 Dissemination ---
    add_heading_styled(doc, "2.2 Communication, Dissemination, and Exploitation", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The dissemination strategy leverages the platform's open-source nature and web-based accessibility:")
    
    dissem = [
        ("Open-Source Repository: ", "Full codebase published under an open licence, enabling community "
         "contributions, academic forks, and regulatory adaptation."),
        ("Public Web Platform: ", "The deployed application at sustainabilitysignals.com provides free access "
         "to the coverage universe, DQ scoring, and entity extraction tools."),
        ("Academic Publications: ", "Peer-reviewed papers on the DQ scoring methodology, entity extraction "
         "pipeline, and comparative ESG disclosure quality studies across sectors and jurisdictions."),
        ("API Access: ", "RESTful API endpoints (disclosure-quality, entity-extract, chat) available for "
         "third-party integration, enabling ecosystem development."),
        ("Regulatory Engagement: ", "Direct engagement with EFRAG, national competent authorities, and the "
         "European Commission's Sustainable Finance Platform to inform disclosure quality standards."),
    ]
    for prefix, text in dissem:
        add_bullet(doc, text, bold_prefix=prefix, font_size=11)

    # --- 2.3 Sustainability ---
    add_heading_styled(doc, "2.3 Sustainability and Long-Term Viability", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The platform's sustainability is ensured through: (a) serverless, pay-per-use infrastructure with no "
        "fixed server costs; (b) edge-native R2 storage with long-lived immutable caching; (c) deterministic "
        "scoring that requires no model retraining; (d) modular architecture enabling incremental feature "
        "additions without architectural changes; and (e) version-controlled score migration ensuring backward "
        "compatibility as the methodology evolves (current: regex-v4.1, with automatic migration from older versions).")

    # =========================================================================
    # SECTION 3: IMPLEMENTATION
    # =========================================================================
    add_heading_styled(doc, "3. Implementation", level=1, font_size=15,
                       color=RGBColor(0x1B, 0x5E, 0x20), space_before=Pt(14))
    
    # --- 3.1 Work Plan ---
    add_heading_styled(doc, "3.1 Work Plan and Work Packages", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc, "The project is structured into five work packages over 36 months:")
    
    create_table(doc,
        ["WP", "Title", "Lead", "Months", "Key Deliverables"],
        [
            ["WP1", "Platform Core & DQ Engine", "Technical Lead", "M1–M12", "DQ regex-v5.0 engine, enhanced subscore models, ESRS datapoint mapping"],
            ["WP2", "Entity Extraction & ML", "ML Engineer", "M6–M24", "FinBERT-ESG v2 (fine-tuned), 15+ entity classes, sector-specific extractors"],
            ["WP3", "RAG & Knowledge Graph", "AI Engineer", "M12–M30", "Enhanced RAG with cross-report retrieval, ESG knowledge graph, citation chains"],
            ["WP4", "Coverage & Validation", "Data Lead", "M1–M36", "1,000+ reports indexed, cross-sector validation study, benchmark dataset"],
            ["WP5", "Dissemination & Impact", "Project Coordinator", "M1–M36", "2 journal papers, 3 conference presentations, regulatory engagement report"],
        ],
        col_widths=[Cm(1.2), Cm(3.5), Cm(2.5), Cm(1.8), Cm(8.0)]
    )
    
    add_paragraph(doc, "", space_after=Pt(4))
    add_paragraph(doc, "3.1.1 Implementation Timeline", bold=True, font_size=11, space_after=Pt(3))
    
    create_table(doc,
        ["Phase", "Period", "Activities"],
        [
            ["Phase 1: Foundation", "M1–M12", "DQ engine v5.0 development, expanded feature families, "
             "ESRS datapoint registry, initial validation on 200+ reports, frontend UX enhancements"],
            ["Phase 2: Intelligence", "M13–M24", "FinBERT-ESG v2 fine-tuning, expanded entity classes, "
             "cross-report RAG, knowledge graph prototype, 500+ reports indexed"],
            ["Phase 3: Scale & Impact", "M25–M36", "Full-scale validation study, regulatory engagement, "
             "1,000+ reports, academic publications, open benchmark dataset release"],
        ],
        col_widths=[Cm(3.0), Cm(2.0), Cm(12.0)]
    )

    # --- 3.2 Management ---
    add_heading_styled(doc, "3.2 Management Structure and Procedures", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "Project management follows an agile methodology with two-week sprint cycles, continuous integration "
        "and deployment (CI/CD) via GitHub Actions and Cloudflare Pages, and automated quality gates (ESLint, "
        "TypeScript strict mode). Code review is mandatory for all changes to scoring logic. The project "
        "coordinator oversees milestone tracking, risk management, and stakeholder communication.")
    
    add_paragraph(doc,
        "Quality assurance for the DQ scoring engine includes: (a) regression tests against a validated corpus; "
        "(b) version-controlled method migration with automatic score recomputation; (c) evidence quote validation "
        "ensuring numerical fidelity; and (d) cross-validation against manually assessed reports.")

    # --- 3.3 Resources ---
    add_heading_styled(doc, "3.3 Resources and Budget Overview", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The project leverages primarily serverless and edge-computing infrastructure, significantly reducing "
        "operational costs compared to traditional cloud deployments:")
    
    create_table(doc,
        ["Resource Category", "Description", "Estimated Annual Cost"],
        [
            ["Compute (Edge)", "Cloudflare Workers/Pages Functions (serverless)", "€2,400/year"],
            ["Storage (R2)", "PDF storage, cached markdown, DQ score JSON", "€1,200/year"],
            ["AI Inference", "Workers AI (chat, markdown, refinement)", "€3,600/year"],
            ["ML Infrastructure", "Azure Container Apps (FinBERT-ESG inference)", "€4,800/year"],
            ["Vectorize", "Semantic search index for RAG", "€600/year"],
            ["Personnel", "2 FTE (Technical Lead + ML Engineer)", "Cost-neutral (in-kind)"],
        ],
        col_widths=[Cm(3.5), Cm(8.5), Cm(4.5)]
    )
    
    add_paragraph(doc, "", space_after=Pt(6))
    
    # =========================================================================
    # SECTION 4: TECHNICAL SPECIFICATIONS (Additional)
    # =========================================================================
    add_heading_styled(doc, "4. API and Data Specifications", level=1, font_size=15,
                       color=RGBColor(0x1B, 0x5E, 0x20), space_before=Pt(14))
    
    add_heading_styled(doc, "4.1 API Endpoints", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    create_table(doc,
        ["Endpoint", "Method", "Function", "Auth"],
        [
            ["/r2/reports/*", "GET, HEAD", "PDF byte serving with Range support, immutable caching", "Public"],
            ["/chat", "POST", "Report-grounded Q&A (RAG), background Vectorize ingestion", "Rate-limited"],
            ["/score/disclosure-quality", "GET", "Cached DQ score retrieval, optional refinement & migration", "Public"],
            ["/score/disclosure-quality", "POST", "DQ score computation: PDF→Markdown→Features→Score→Cache", "Rate-limited"],
            ["/score/disclosure-quality-batch", "POST", "Bulk DQ summary fetch (up to 200 IDs)", "Rate-limited"],
            ["/score/entity-extract", "GET/POST", "Hybrid FinBERT + LLM entity extraction pipeline", "Rate-limited"],
            ["/api/reports/upload", "POST", "PDF upload with AI metadata suggestion, duplicate detection", "Authenticated"],
        ],
        col_widths=[Cm(4.5), Cm(2.0), Cm(8.0), Cm(2.5)]
    )
    
    add_paragraph(doc, "", space_after=Pt(3))
    add_heading_styled(doc, "4.2 Data Model and Output Contract", level=2, font_size=13,
                       color=RGBColor(0x2E, 0x7D, 0x32))
    
    add_paragraph(doc,
        "The DQ scoring output includes: version, generatedAt, report metadata (id, key, company, year), "
        "overall score (0–100) and band (high/medium/low), four subscores, boolean feature map with evidence "
        "quotes (ranked by quality), quantitative profile (percentage count, table rows, KPI numbers, distinct "
        "years, numeric density), improvement recommendations (top 6), and method metadata (kind, weights, "
        "corpus statistics, migration history). All results are cached in R2 at "
        "scores/disclosure_quality/v{version}/{reportId}.json with automatic version migration.")
    
    add_paragraph(doc,
        "The entity extraction output includes: extracted entities (name, value, unit, class, pillar, source chunk), "
        "FinBERT routing statistics (ESG vs. Non-ESG chunk ratios), and per-entity confidence indicators from "
        "the LLM extraction stage.")

    # =========================================================================
    # SECTION 5: ETHICAL AND SOCIETAL CONSIDERATIONS
    # =========================================================================
    add_heading_styled(doc, "5. Ethical and Societal Considerations", level=1, font_size=15,
                       color=RGBColor(0x1B, 0x5E, 0x20), space_before=Pt(14))
    
    add_paragraph(doc,
        "The platform operates under the following ethical principles: (a) Transparency — all scoring algorithms "
        "are open-source and fully documented; (b) Reproducibility — deterministic regex-based scoring ensures "
        "identical results for identical inputs; (c) Disclosure — the platform explicitly states that it provides "
        "disclosure quality assessment, not investment advice; (d) Bias Mitigation — regex-based detection avoids "
        "the implicit biases of ML-trained scoring models; (e) Data Minimisation — no personal data is collected "
        "or processed; the platform analyses only publicly available corporate sustainability reports.")
    
    add_paragraph(doc,
        "The project complies with the EU AI Act by maintaining human oversight (manual score review capability), "
        "providing full algorithmic transparency, and documenting known limitations (regex sensitivity to "
        "paraphrasing, OCR quality dependency, sampling effects on large documents).")

    # =========================================================================
    # SAVE
    # =========================================================================
    out_path = os.path.abspath(OUTPUT)
    doc.save(out_path)
    print(f"Technical Description saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    build_document()
